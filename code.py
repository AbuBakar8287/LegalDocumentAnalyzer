import re
import os
from dotenv import load_dotenv
import fitz
import docx
from langchain_community.vectorstores import FAISS

# LangChain components
from langchain_community.vectorstores import chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.llms import HuggingFacePipeline
# from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.embeddings import HuggingFaceEmbeddings

from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnableParallel, RunnablePassthrough, RunnableLambda
from langchain_core.output_parsers import StrOutputParser
from langchain.retrievers import ContextualCompressionRetriever
from langchain.retrievers.document_compressors import LLMChainFilter

# Hugging Face Transformers
from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM


# Cohere for reranking
import cohere

# Load environment variables
load_dotenv()
os.environ["COHERE_API_KEY"] = "ZNfhw5DTkPxmRK9vNON7gc4zrQQ1WfrcctNkMdUl3"
co = cohere.Client(os.getenv("COHERE_API_KEY"))
os.environ["HUGGINGFACEHUB_API_TOKEN"] = "hf_vyjWJKxdOvraJwwurOlQjDBkevrKxsaxf"
def extract_text_from_pdf(pdf_path):
    text = ""
    doc = fitz.open(pdf_path)
    for page in doc:
        text += page.get_text()
    return text


def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])
import cv2
import pytesseract
from PIL import Image
import numpy as np

# Optional: Configure Tesseract path if needed
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

def preprocess_image(image_path):
    """Preprocess the image to improve OCR accuracy."""

    # Read image in grayscale
    img = cv2.imread(image_path, cv2.IMREAD_GRAYSCALE)

    # Denoising
    img = cv2.medianBlur(img, 3)

    # Thresholding (Binarization)
    _, img = cv2.threshold(img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)

    # (Optional) Resize for better OCR if image is small
    scale_percent = 150  # scale up by 150%
    width = int(img.shape[1] * scale_percent / 100)
    height = int(img.shape[0] * scale_percent / 100)
    img = cv2.resize(img, (width, height), interpolation=cv2.INTER_LINEAR)

    return img

def ocr_image(preprocessed_img):
    """Apply OCR to preprocessed image."""

    text = pytesseract.image_to_string(preprocessed_img)
    return text
  from PIL import Image
import pytesseract
def extract_text_from_image(file_path):
    image = Image.open(file_path)
    text = pytesseract.image_to_string(image)
    return text
  def extract_text(file_path, file_type):
    file_type = file_type.lower()

    if file_type == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type == "docx":
        return extract_text_from_docx(file_path)
    elif file_type == "txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    elif file_type in ["jpg", "jpeg", "png"]:
        return extract_text_from_image(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
  # Example: After extracting text
text = extract_text("/content/free-software-support-agreement.docx", "docx")
from nltk.tokenize import sent_tokenize
# Tokenize text into sentences
sentences = sent_tokenize(text)

# Parameters for semantic chunking
chunk_size = 1000  # target number of characters per chunk
chunk_overlap = 200

# Create semantic chunks
chunks = []
current_chunk = ""

for sentence in sentences:
    if len(current_chunk) + len(sentence) <= chunk_size:
        current_chunk += " " + sentence
    else:
        chunks.append(current_chunk.strip())

        # Add overlap: take last chunk_overlap characters from the previous chunk
        overlap = current_chunk[-chunk_overlap:] if chunk_overlap > 0 else ""

        current_chunk = overlap + " " + sentence

# Append any remaining text
if current_chunk:
    chunks.append(current_chunk.strip())

print(f"Total Chunks: {len(chunks)}")
#embeddings
from langchain.embeddings import huggingface
from langchain.embeddings import HuggingFaceEmbeddings
embedding = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
#vector store
from langchain.docstore.document import Document
documents = [Document(page_content=item) for item in chunks]
vector_store = FAISS.from_documents(documents, embedding)
#retriever
retriever = vector_store.as_retriever(search_type="mmr", search_kwargs={"k": 4, "lambda_mult": 0.5})
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

model_id = "microsoft/phi-2"

# Force fresh download to fix corrupted tokenizer file
tokenizer = AutoTokenizer.from_pretrained(model_id, force_download=True)
model = AutoModelForCausalLM.from_pretrained(model_id, force_download=True)

pipe = pipeline("text-generation", model=model, tokenizer=tokenizer, max_new_tokens=100, temperature=0.3, do_sample=True)
llm = HuggingFacePipeline(pipeline=pipe)
from langchain_core.prompts import PromptTemplate
import re
from langchain_core.runnables import RunnableParallel, RunnablePassthrough, RunnableLambda ,RunnableMap
from langchain_core.output_parsers import StrOutputParser

# Prompt Template
prompt_template = """
You are a Legal Document Assistant. Based on the extracted legal context below, answer the user's question concisely and accurately.
If applicable, include the clause number (e.g., Clause 4.4) where the answer is found.

context:{context}

question: {question}

If the answer is not found in the context, say "I don't know based on the provided document."
"""

prompt = PromptTemplate.from_template(prompt_template)
# Clause extractor
def extract_clause_references(text):
    pattern = r"(Clause\s\d+(\.\d+)?)"
    matches = re.findall(pattern, text)
    return list(set(match[0] for match in matches))

# Final answer formatter
def postprocess_with_clauses(answer, context):
    clause_refs = extract_clause_references(context)
    if clause_refs:
        return f"{answer}\n\nReferenced {', '.join(clause_refs)}"
    return answer

from langchain_core.runnables import Runnable

class CohereRerankRetriever(Runnable):
    def __init__(self, retriever, cohere_client, top_n=4):
        self.retriever = retriever
        self.cohere_client = cohere_client
        self.top_n = top_n

    def invoke(self, query, config=None):
        raw_docs = self.retriever.invoke(query)
        candidates = [doc.page_content for doc in raw_docs]
        response = self.cohere_client.rerank(query=query, documents=candidates, top_n=self.top_n)
        reranked_docs = [raw_docs[result.index] for result in response.results]
        return reranked_docs
# ✅ Ensure you have a working retriever and cohere client
rerank_retriever = CohereRerankRetriever(retriever, cohere_client=co, top_n=2)
def format_docs(docs):
    return "\n\n".join(doc.page_content for doc in docs)

# RAG Chain
base_rag_chain = (
    {"context": rerank_retriever| RunnableLambda(format_docs), "question": RunnablePassthrough()}
    | prompt
    | llm
    | StrOutputParser()

)
# Step 2: Wrap input and keep full context + answer for clause postprocessing
final_chain = (
    {
        "context": rerank_retriever | RunnableLambda(format_docs),
        "question": RunnablePassthrough()
    }
    | RunnableMap({
        "context": lambda x: x["context"],
        "question": lambda x: x["question"],
        "answer": lambda x: base_rag_chain.invoke(x["question"])
    })
    | RunnableLambda(lambda x: postprocess_with_clauses(x["answer"], x["context"]))
)


# Test RAG
question = "Are there any late payment penalties ?"
response = final_chain.invoke(question)
print(response)
